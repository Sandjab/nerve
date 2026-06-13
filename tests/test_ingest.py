import io
import zipfile
import pytest
import nerve.ingest as ing
from nerve.ingest import IngestError

def test_read_txt(tmp_path):
    p = tmp_path / "a.txt"
    p.write_text("contenu texte", encoding="utf-8")
    assert ing.read_file(str(p), "a.txt") == "contenu texte"

def test_read_empty_raises(tmp_path):
    p = tmp_path / "v.txt"
    p.write_text("   \n  ", encoding="utf-8")
    with pytest.raises(IngestError):
        ing.read_file(str(p), "v.txt")

def test_read_docx(tmp_path):
    from docx import Document
    d = Document()
    d.add_paragraph("Bonjour")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "A"
    t.rows[0].cells[1].text = "B"
    p = tmp_path / "doc.docx"
    d.save(str(p))
    out = ing.read_file(str(p), "doc.docx")
    assert "Bonjour" in out
    assert "A" in out and "B" in out

def test_read_pdf_via_pypdf(tmp_path, monkeypatch):
    class _Page:
        def extract_text(self):
            return "page un"
    class _Reader:
        def __init__(self, path):
            self.pages = [_Page()]
    monkeypatch.setattr(ing, "PdfReader", _Reader)
    p = tmp_path / "x.pdf"
    p.write_bytes(b"%PDF-1.4 factice")
    assert ing.read_file(str(p), "x.pdf") == "page un"

def test_read_pdf_corrupt_raises(tmp_path, monkeypatch):
    def _boom(path):
        raise ValueError("pdf cassé")
    monkeypatch.setattr(ing, "PdfReader", _boom)
    p = tmp_path / "bad.pdf"
    p.write_bytes(b"pas un pdf")
    with pytest.raises(IngestError):
        ing.read_file(str(p), "bad.pdf")

def test_read_html_via_trafilatura(tmp_path, monkeypatch):
    monkeypatch.setattr(ing.trafilatura, "extract", lambda raw, **k: "# T\n\ncorps")
    p = tmp_path / "page.html"
    p.write_text("<html><body><p>x</p></body></html>", encoding="utf-8")
    assert "corps" in ing.read_file(str(p), "page.html")


def _make_zip(entries: dict) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        for name, content in entries.items():
            z.writestr(name, content)
    return buf.getvalue()

def test_ingest_zip_resilient(tmp_path):
    raw = _make_zip({"a.txt": "contenu A", "empty.txt": "", "../evil.txt": "x"})
    segments, skipped = ing.ingest_upload("c.zip", raw, str(tmp_path / "dest"))
    assert segments == [("contenu A", "a.txt")]      # empty ignoré, evil non extrait
    assert skipped == ["empty.txt"]

def test_ingest_zip_all_unreadable_raises(tmp_path):
    raw = _make_zip({"empty.txt": "   "})
    with pytest.raises(IngestError):
        ing.ingest_upload("c.zip", raw, str(tmp_path / "dest"))

def test_ingest_single_file(tmp_path):
    segments, skipped = ing.ingest_upload("note.txt", b"un seul fichier", str(tmp_path / "d2"))
    assert segments == [("un seul fichier", "")]
    assert skipped == []

def test_ingest_single_file_empty_raises(tmp_path):
    with pytest.raises(IngestError):
        ing.ingest_upload("vide.txt", b"   ", str(tmp_path / "d3"))

def test_read_docx_merged_cells_no_dup(tmp_path):
    # python-docx répète une cellule fusionnée dans row.cells -> ne pas dupliquer son texte
    from docx import Document
    d = Document()
    t = d.add_table(rows=1, cols=3)
    t.cell(0, 0).text = "X"; t.cell(0, 1).text = "Y"; t.cell(0, 2).text = "Z"
    t.cell(0, 0).merge(t.cell(0, 1))      # X et Y fusionnent (texte "X\nY", répété dans row.cells)
    p = tmp_path / "m.docx"; d.save(str(p))
    out = ing.read_file(str(p), "m.docx")
    assert out.count("X") == 1            # le texte fusionné n'apparaît qu'une fois
    assert "Z" in out

def test_ingest_corrupt_zip_raises(tmp_path):
    # un .zip qui n'est pas une archive valide -> IngestError (pas BadZipFile nue)
    with pytest.raises(IngestError):
        ing.ingest_upload("c.zip", b"ceci n'est pas un zip", str(tmp_path / "dz"))

def test_ingest_zip_skips_non_ingesterror(tmp_path, monkeypatch):
    # une erreur de lecture NON-IngestError (ici trafilatura qui lève) ne doit pas
    # avorter le lot : le fichier fautif part dans skipped, les autres survivent.
    def _boom(raw, **k):
        raise ValueError("html cassé")
    monkeypatch.setattr(ing.trafilatura, "extract", _boom)
    raw = _make_zip({"a.txt": "bon contenu", "b.html": "<html><body>x</body></html>"})
    segments, skipped = ing.ingest_upload("c.zip", raw, str(tmp_path / "dz2"))
    assert segments == [("bon contenu", "a.txt")]
    assert skipped == ["b.html"]
