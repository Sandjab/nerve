# nerve/ingest.py
import os
import shutil
import zipfile
import trafilatura
from pypdf import PdfReader
from docx import Document


class IngestError(Exception):
    """Lecture impossible (fichier vide, illisible, format cassé)."""


TEXT_EXTS = {".txt", ".md", ".markdown", ".text", ".rst", ".json", ".jsonl",
             ".csv", ".tsv", ".log", ".html", ".htm", ".pdf", ".docx",
             ".xml", ".yaml", ".yml", ".py", ".js", ".ts"}


def _read_pdf(path: str) -> str:
    try:
        reader = PdfReader(path)
        return "\n\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception as e:
        raise IngestError(f"PDF illisible : {e}")


def _read_docx(path: str) -> str:
    try:
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception as e:
        raise IngestError(f"docx illisible : {e}")


def _read_html(path: str) -> str:
    try:
        with open(path, encoding="utf-8", errors="ignore") as f:
            raw = f.read()
        return trafilatura.extract(raw, output_format="markdown") or ""
    except Exception as e:
        raise IngestError(f"HTML illisible : {e}")


def _read_text(path: str) -> str:
    try:
        with open(path, encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        raise IngestError(f"Texte illisible : {e}")


def read_file(path: str, name: str) -> str:
    """Extrait le texte d'un fichier. Lève IngestError si vide/illisible.
    Toutes les branches convertissent leurs erreurs en IngestError -> la boucle
    zip (qui n'attrape qu'IngestError) reste résiliente quel que soit le format."""
    ext = os.path.splitext(name)[1].lower()
    if ext == ".pdf":
        text = _read_pdf(path)
    elif ext == ".docx":
        text = _read_docx(path)
    elif ext in (".html", ".htm"):
        text = _read_html(path)
    else:
        text = _read_text(path)
    if not text or not text.strip():
        raise IngestError(f"Contenu vide/illisible : {name}")
    return text


def ingest_upload(filename: str, raw: bytes, dest_dir: str) -> tuple[list[tuple[str, str]], list[str]]:
    """Ingestion d'un upload. Conserve le brut sous dest_dir.
    Retourne (segments[(text, source_file)], skipped[noms ignorés pour illisibilité]).
    Zip : résilient (un fichier illisible -> skipped). Mono-fichier ou zip 100 %
    illisible -> IngestError (fail-loud)."""
    os.makedirs(dest_dir, exist_ok=True)
    raw_path = os.path.join(dest_dir, os.path.basename(filename))
    with open(raw_path, "wb") as f:
        f.write(raw)
    if filename.lower().endswith(".zip"):
        return _ingest_zip(raw_path, dest_dir)
    text = read_file(raw_path, filename)        # IngestError remonte (fail-loud)
    return [(text, "")], []


def _ingest_zip(zip_path: str, dest_dir: str) -> tuple[list[tuple[str, str]], list[str]]:
    segments: list[tuple[str, str]] = []
    skipped: list[str] = []
    had_supported = False
    root = os.path.abspath(dest_dir) + os.sep
    try:
        zf = zipfile.ZipFile(zip_path)
    except (zipfile.BadZipFile, OSError) as e:
        raise IngestError(f"Zip illisible : {e}")
    with zf:
        for info in sorted(zf.infolist(), key=lambda i: i.filename):
            if info.is_dir():
                continue
            name = info.filename
            base = os.path.basename(name)
            if not base or base.startswith(".") or "__MACOSX" in name:
                continue
            if os.path.splitext(base)[1].lower() not in TEXT_EXTS:
                continue
            target = os.path.join(dest_dir, name)
            if not os.path.abspath(target).startswith(root):   # anti path-traversal
                continue
            had_supported = True
            os.makedirs(os.path.dirname(target), exist_ok=True)
            with zf.open(info) as src, open(target, "wb") as dst:
                shutil.copyfileobj(src, dst)
            try:
                text = read_file(target, name)
            except IngestError:
                skipped.append(name)
                continue
            segments.append((text, name))
    if had_supported and not segments:
        raise IngestError(f"Aucun fichier lisible dans {os.path.basename(zip_path)}")
    return segments, skipped
